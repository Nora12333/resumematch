import json
import logging
import re
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import Response
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models import GenerateRequest, GenerateResponse
from services.llm import generate_resume, generate_resume_structured
from services.matcher import ResumeMatcher

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api", tags=["generation"])


@router.post("/generate", response_model=GenerateResponse)
def generate_resume_endpoint(payload: GenerateRequest) -> GenerateResponse:
    gaps_dicts = [g.model_dump() for g in payload.gaps]
    try:
        result = generate_resume(
            resume_text=payload.resume_text,
            jd_text=payload.jd_text,
            gaps=gaps_dicts,
            mode=payload.mode,
        )
        try:
            matcher = ResumeMatcher()
            score_after = matcher.analyze(
                resume_text=result["optimized_resume"],
                jd_text=payload.jd_text,
            )
            result["score_after"] = score_after
        except Exception as e:
            logger.warning(f"Re-scoring failed: {e}")
            result["score_after"] = None
        # Also generate structured data for docx download
        try:
            structured = generate_resume_structured(
                resume_text=payload.resume_text,
                jd_text=payload.jd_text,
                gaps=gaps_dicts,
                mode=payload.mode,
            )
            result["structured"] = structured
        except Exception as e:
            logger.warning(f"Structured generation failed: {e}")
            result["structured"] = None
        return GenerateResponse(**result)
    except RuntimeError as e:
        msg = str(e)
        if "ANTHROPIC_API_KEY" in msg or "not installed" in msg:
            raise HTTPException(status_code=503, detail=msg) from e
        if msg.startswith("Claude API error:"):
            raise HTTPException(status_code=502, detail=msg) from e
        raise HTTPException(status_code=500, detail=msg) from e
    except (json.JSONDecodeError, ValueError) as e:
        raise HTTPException(status_code=502, detail=f"Invalid response from language model: {e}") from e
    except Exception as e:
        raise HTTPException(status_code=502, detail=str(e)) from e


def _parse_segments(text):
    result = []
    remaining = text
    while '[NEW]' in remaining:
        start = remaining.find('[NEW]')
        if start > 0:
            result.append((remaining[:start], False))
        remaining = remaining[start+5:]
        end = remaining.find('[NEW]')
        if end == -1:
            result.append((remaining, False))
            remaining = ''
            break
        result.append((remaining[:end], True))
        remaining = remaining[end+5:]
    if remaining:
        result.append((remaining, False))
    return result


def _add_runs(para, text, bold=False, italic=False, size=10):
    for seg, is_new in _parse_segments(text):
        if not seg:
            continue
        run = para.add_run(seg)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Arial"


def _add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_right_tab(p, content_w=9360):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(content_w))
    tabs.append(tab)
    pPr.append(tabs)


def _build_docx_structured(data: dict, pages: int = 1) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    margin = Inches(0.6) if pages == 1 else Inches(0.75)
    section.top_margin = section.bottom_margin = margin
    section.left_margin = section.right_margin = margin

    body_size = 9 if pages == 1 else 10
    name_size = 14 if pages == 1 else 16
    sec_size = 10 if pages == 1 else 11
    sp = 1 if pages == 1 else 3
    content_w = 9360 if pages == 2 else 9072

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, data.get("name", ""), bold=True, size=name_size)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sp)
    _add_runs(p, data.get("contact", ""), size=body_size)

    # EDUCATION
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    _add_runs(p, "EDUCATION", bold=True, size=sec_size)
    _add_hr(doc)

    for edu in data.get("education", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(edu.get("school", ""))
        run.bold = True
        run.font.size = Pt(body_size)
        run.font.name = "Arial"

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _add_right_tab(p, content_w)
        run = p.add_run(edu.get("degree", ""))
        run.italic = True
        run.font.size = Pt(body_size)
        run.font.name = "Arial"
        run2 = p.add_run("\t" + edu.get("date", ""))
        run2.italic = True
        run2.font.size = Pt(body_size)
        run2.font.name = "Arial"

        if edu.get("coursework"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(sp)
            r = p.add_run("Relevant Coursework: ")
            r.italic = True
            r.font.size = Pt(body_size)
            r.font.name = "Arial"
            _add_runs(p, edu["coursework"], size=body_size)

    # EXPERIENCE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    _add_runs(p, "EXPERIENCE", bold=True, size=sec_size)
    _add_hr(doc)

    for exp in data.get("experience", []):
        org = exp.get("company", "")
        if exp.get("location"):
            org += f", {exp['location']}"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _add_runs(p, org, bold=True, size=body_size)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sp)
        _add_right_tab(p, content_w)
        run = p.add_run(exp.get("title", ""))
        run.italic = True
        run.font.size = Pt(body_size)
        run.font.name = "Arial"
        run2 = p.add_run("\t" + exp.get("date", ""))
        run2.italic = True
        run2.font.size = Pt(body_size)
        run2.font.name = "Arial"

        for bullet in exp.get("bullets", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(sp)
            p.paragraph_format.left_indent = Inches(0.15)
            run = p.add_run("• ")
            run.font.size = Pt(body_size)
            run.font.name = "Arial"
            _add_runs(p, bullet, size=body_size)

    # PROJECTS
    if data.get("projects"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        _add_runs(p, "PROJECTS", bold=True, size=sec_size)
        _add_hr(doc)

        for proj in data.get("projects", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(sp)
            _add_runs(p, proj.get("name", ""), bold=True, size=body_size)
            for bullet in proj.get("bullets", []):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(sp)
                p.paragraph_format.left_indent = Inches(0.15)
                run = p.add_run("• ")
                run.font.size = Pt(body_size)
                run.font.name = "Arial"
                _add_runs(p, bullet, size=body_size)

    # SKILLS
    if data.get("skills"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        _add_runs(p, "SKILLS", bold=True, size=sec_size)
        _add_hr(doc)

        for skill in data.get("skills", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(sp)
            p.paragraph_format.left_indent = Inches(0.15)
            run = p.add_run("• ")
            run.font.size = Pt(body_size)
            run.font.name = "Arial"
            r = p.add_run(skill.get("label", "") + ":  ")
            r.bold = True
            r.font.size = Pt(body_size)
            r.font.name = "Arial"
            _add_runs(p, skill.get("content", ""), size=body_size)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


from pydantic import BaseModel as PM

class DocxFromTextRequest(PM):
    optimized_text: str
    pages: int = 2


@router.post("/generate-docx")
async def generate_docx_endpoint(request: Request, pages: int = 2):
    body = await request.json()
    optimized_text = body.get("optimized_text", "") or body.get("resume_text", "")
    clean_text = optimized_text.replace("[NEW]", "").strip()

    SECTION_HEADERS = {"EDUCATION", "EXPERIENCE", "PROJECTS", "SKILLS", "SUMMARY", "CERTIFICATIONS"}

    try:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        margin = Inches(0.75) if pages >= 2 else Inches(0.6)
        section.top_margin = section.bottom_margin = margin
        section.left_margin = section.right_margin = margin

        lines = clean_text.split("\n")
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                continue

            if idx == 0:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = "Arial"

            elif idx == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.name = "Arial"

            elif stripped.upper() in SECTION_HEADERS:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(stripped.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bot = OxmlElement("w:bottom")
                bot.set(qn("w:val"), "single")
                bot.set(qn("w:sz"), "6")
                bot.set(qn("w:space"), "1")
                bot.set(qn("w:color"), "000000")
                pBdr.append(bot)
                pPr.append(pBdr)

            elif stripped.startswith("•"):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Inches(0.15)
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.name = "Arial"

            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.name = "Arial"

        buf = BytesIO()
        doc.save(buf)
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.post("/generate-docx-text")
def generate_docx_from_text(payload: DocxFromTextRequest):
    try:
        # Convert plain text to simple docx
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        margin = Inches(0.75) if payload.pages >= 2 else Inches(0.6)
        section.top_margin = section.bottom_margin = margin
        section.left_margin = section.right_margin = margin
        
        for line in payload.optimized_text.split("\n"):
            clean = line.replace("[NEW]", "").strip()
            if not clean:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean)
            run.font.name = "Arial"
            run.font.size = Pt(10)
        
        buf = BytesIO()
        doc.save(buf)
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e

from pydantic import BaseModel
class StructuredDocxRequest(BaseModel):
    structured: dict
    pages: int = 1

@router.post("/download-docx")  
def download_docx_from_structured(payload: StructuredDocxRequest):
    try:
        docx_bytes = _build_docx_structured(payload.structured, pages=payload.pages)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e
