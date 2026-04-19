import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HIGHLIGHT_COLOR = "E6FFE6"  # 浅绿色高亮

def _set_cell_highlight(run, color_hex: str):
    """给 run 加背景高亮色"""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'green')
    rPr.append(highlight)

def _add_hr(doc):
    """添加横线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def _parse_segments(text: str):
    """把文本按 [NEW]...[NEW] 拆成 [(text, is_new), ...]"""
    parts = re.split(r'(\[NEW\][\s\S]*?\[NEW\])', text)
    segments = []
    for part in parts:
        if re.match(r'^\[NEW\][\s\S]*\[NEW\]$', part):
            cleaned = part[5:-5]  # 去掉首尾 [NEW]
            segments.append((cleaned, True))
        elif part:
            segments.append((part, False))
    return segments

def _add_mixed_run(para, text: str, bold=False, italic=False, size_pt=10):
    """在段落里添加含高亮的混合文本"""
    segments = _parse_segments(text)
    for seg_text, is_new in segments:
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = 'Arial'
        if is_new:
            # 黄色高亮
            rPr = run._r.get_or_add_rPr()
            hl = OxmlElement('w:highlight')
            hl.set(qn('w:val'), 'yellow')
            rPr.append(hl)

def build_resume_docx(optimized_resume: str, pages: int = 2) -> bytes:
    """
    把 optimized_resume 文本生成格式化 docx。
    pages: 1 或 2，控制字号和间距
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    if pages == 1:
        margin = Inches(0.6)
    else:
        margin = Inches(0.75)
    section.top_margin = section.bottom_margin = margin
    section.left_margin = section.right_margin = margin

    # 字号配置
    name_size   = 14 if pages == 1 else 16
    body_size   = 9  if pages == 1 else 10
    sec_size    = 10 if pages == 1 else 11
    sp_after_p  = 1  if pages == 1 else 3   # 段落间距 pt
    sp_sec      = 4  if pages == 1 else 6

    # 清除默认样式的多余间距
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(body_size)

    lines = optimized_resume.split('\n')
    i = 0

    def is_section_header(line):
        stripped = line.strip()
        keywords = ['EDUCATION', 'EXPERIENCE', 'PROJECTS', 'SKILLS', 'SUMMARY',
                    '教育', '经历', '技能', '项目']
        return stripped.upper() in [k.upper() for k in keywords] or \
               (stripped.isupper() and 3 < len(stripped) < 30 and not stripped.startswith('•'))

    def is_bullet(line):
        return line.strip().startswith('•') or line.strip().startswith('-')

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        # 姓名（第一行非空）
        if i == 0 or (i < 3 and not any(c in stripped for c in ['|', '@', '+'])):
            # 找真正的姓名行
            if not any(c in stripped for c in ['|', '@', '+']):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(1)
                _add_mixed_run(p, stripped, bold=True, size_pt=name_size)
                i += 1
                continue

        # 联系信息行（含 | 或 @）
        if '|' in stripped or ('@' in stripped and 'edu' not in stripped.lower()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(sp_after_p)
            _add_mixed_run(p, stripped, size_pt=body_size)
            i += 1
            continue

        # Section header
        if is_section_header(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(sp_sec)
            p.paragraph_format.space_after = Pt(0)
            _add_mixed_run(p, stripped, bold=True, size_pt=sec_size)
            _add_hr(doc)
            i += 1
            continue

        # Bullet point
        if is_bullet(stripped):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(sp_after_p)
            p.paragraph_format.left_indent = Inches(0.2)
            # 去掉自带bullet符号，用文字里的•
            bullet_text = re.sub(r'^[•\-]\s*', '', stripped)
            _add_mixed_run(p, bullet_text, size_pt=body_size)
            i += 1
            continue

        # Relevant Coursework 行
        if stripped.startswith('Relevant Coursework'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(sp_after_p)
            if ':' in stripped:
                label, rest = stripped.split(':', 1)
                run1 = p.add_run(label + ':')
                run1.italic = True
                run1.font.size = Pt(body_size)
                run1.font.name = 'Arial'
                _add_mixed_run(p, rest, size_pt=body_size)
            else:
                _add_mixed_run(p, stripped, italic=True, size_pt=body_size)
            i += 1
            continue

        # 时间范围行（含 – 或 -，且不是bullet）
        if ('–' in stripped or ' - ' in stripped) and not is_bullet(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(sp_after_p)
            # 尝试左右对齐
            if '\t' in stripped:
                left, right = stripped.split('\t', 1)
                _add_mixed_run(p, left, italic=True, size_pt=body_size)
                run_tab = p.add_run('\t')
                run_tab.font.size = Pt(body_size)
                _add_mixed_run(p, right.strip(), italic=True, size_pt=body_size)
                from docx.oxml import OxmlElement
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9360')
                tabs.append(tab)
                pPr.append(tabs)
            else:
                _add_mixed_run(p, stripped, italic=True, size_pt=body_size)
            i += 1
            continue

        # 其他普通行（公司名、学校名等）
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sp_after_p)
        # 判断是否是机构/公司名（不含逗号日期则加粗）
        looks_like_org = (
            not stripped.startswith('•') and
            len(stripped) < 80 and
            not re.search(r'\d{4}', stripped)
        )
        _add_mixed_run(p, stripped, bold=looks_like_org, size_pt=body_size)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()